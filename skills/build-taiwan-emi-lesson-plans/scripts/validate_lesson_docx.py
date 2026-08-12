#!/usr/bin/env python3
import argparse, sys
from docx import Document

ap = argparse.ArgumentParser()
ap.add_argument("docx")
ns = ap.parse_args()
d = Document(ns.docx)
text = "\n".join([p.text for p in d.paragraphs] + [c.text for t in d.tables for row in t.rows for c in row.cells])
labels = ["ODIR", "Intended Learning Outcomes", "Disciplinary English", "Teaching Support", "Cognitions"]
checks = {label: label.lower() in text.lower() for label in labels}
for label, ok in checks.items():
    print(f"{label}: {'OK' if ok else 'MISSING'}")
sys.exit(0 if all(checks.values()) else 1)
